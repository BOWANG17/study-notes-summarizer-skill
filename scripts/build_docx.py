#!/usr/bin/env python3
"""
study-notes-summarizer — Markdown 总结 → Word 文档渲染器
========================================================
把 AI 生成的结构化 Markdown 总结（H1/H2/H3、列表、**粗体**、*斜体*、
`代码`、> 引用、表格、⚠️ 易错段）渲染成一份排版干净的 .docx。

设计目标：
  - 只依赖 python-docx（解析 .docx 时已装，无需额外系统软件）。
  - 不依赖 WorkBuddy 内置 tencent-docx，任何能跑 Python 的 AI 工具都能用。
  - ⚠️ 易错段落自动标黄高亮，醒目便于考前避坑。
  - 容错：不认识的行原样作为普通段落保留，不会丢内容。

用法：
  python3 build_docx.py --input summaries/7月总结.md --output final/7月笔记总结.docx [--title "德语B1 7月笔记总结"]
"""
import argparse
import re
import sys
from pathlib import Path


def _check_docx():
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


# ---------- 内联格式 ----------
# 顺序：**bold** → `code` → *italic*
_INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")


def add_inline(paragraph, text):
    """把一行文本按内联格式拆成多个 run 加进段落。"""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
        elif tok.startswith("*"):
            paragraph.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _highlight(paragraph, color="yellow"):
    """给段落所有 run 加底色高亮。"""
    from docx.enum.text import WD_COLOR_INDEX
    mapping = {
        "yellow": WD_COLOR_INDEX.YELLOW,
        "red": WD_COLOR_INDEX.RED,
    }
    hl = mapping.get(color, WD_COLOR_INDEX.YELLOW)
    for r in paragraph.runs:
        r.font.highlight_color = hl


def _set_cjk_font(doc):
    """给 Normal 样式设中文字体，避免 Word 里中文显示异常。"""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        style = doc.styles["Normal"]
        style.font.size = None  # 用默认 11pt
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "宋体")
        rfonts.set(qn("w:ascii"), "Calibri")
        rfonts.set(qn("w:hAnsi"), "Calibri")
    except Exception:
        pass  # 字体设置失败不致命


# ---------- 行级解析 ----------
_BULLET = re.compile(r"^[-*+]\s+(.*)$")
_NUMBER = re.compile(r"^\d+\.\s+(.*)$")
_HR = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")


def render(markdown_text: str, out_path: Path, title_override: str = None):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    _set_cjk_font(doc)

    lines = markdown_text.splitlines()
    title_set = False
    i = 0
    n = len(lines)
    in_warn = False  # 是否处于 ⚠️ 易错小节（整节标黄）

    # 表格缓冲
    table_buf = []

    def flush_table():
        if not table_buf:
            return
        rows = []
        for ln in table_buf:
            cells = [c.strip() for c in ln.strip().strip("|").split("|")]
            rows.append(cells)
        rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
        if rows:
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for ri, r in enumerate(rows):
                for ci in range(ncols):
                    cell = tbl.cell(ri, ci)
                    cell.text = r[ci] if ci < len(r) else ""
            for ci in range(ncols):
                for p in tbl.cell(0, ci).paragraphs:
                    for r in p.runs:
                        r.bold = True
            doc.add_paragraph("")
        table_buf.clear()

    def maybe_warn(p, text):
        if in_warn or "⚠️" in text:
            _highlight(p)

    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("|"):
            table_buf.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        if not stripped:
            i += 1
            continue

        if _HR.match(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 标题：进入新的 H1/H2 时重置 in_warn；标题本身含 ⚠️ 则开启整节高亮
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            in_warn = "⚠️" in text
            if not title_set:
                h = doc.add_heading(text, level=0)
                doc.core_properties.title = text
                title_set = True
            else:
                h = doc.add_heading(text, level=1)
            maybe_warn(h, text)
            i += 1
            continue
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            in_warn = "⚠️" in text
            h = doc.add_heading(text, level=1)
            maybe_warn(h, text)
            i += 1
            continue
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            # H3 不重置 in_warn（易错小节的子标题继续高亮）；自身含 ⚠️ 也开启
            in_warn = in_warn or ("⚠️" in text)
            h = doc.add_heading(text, level=2)
            maybe_warn(h, text)
            i += 1
            continue
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            in_warn = in_warn or ("⚠️" in text)
            h = doc.add_heading(text, level=3)
            maybe_warn(h, text)
            i += 1
            continue

        # 引用块 >
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, text)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            maybe_warn(p, text)
            i += 1
            continue

        # 无序列表
        mb = _BULLET.match(stripped)
        if mb:
            content = mb.group(1)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, content)
            maybe_warn(p, stripped)
            i += 1
            continue

        # 有序列表
        mn = _NUMBER.match(stripped)
        if mn:
            content = mn.group(1)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, content)
            maybe_warn(p, stripped)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_inline(p, stripped)
        maybe_warn(p, stripped)
        i += 1

    flush_table()

    if title_override and not title_set:
        doc.add_heading(title_override, level=0)
        doc.core_properties.title = title_override

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main():
    ap = argparse.ArgumentParser(description="把 Markdown 总结渲染成 Word 文档（python-docx）")
    ap.add_argument("--input", required=True, help="输入的 Markdown 总结文件")
    ap.add_argument("--output", required=True, help="输出的 .docx 路径")
    ap.add_argument("--title", default=None, help="文档标题（无 H1 时使用）")
    args = ap.parse_args()

    if not _check_docx():
        print("❌ 缺少 python-docx。请安装：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    src = Path(args.input).expanduser().resolve()
    if not src.is_file():
        print(f"❌ 输入文件不存在: {src}", file=sys.stderr)
        sys.exit(1)

    text = src.read_text(encoding="utf-8", errors="ignore")
    out = Path(args.output).expanduser().resolve()
    render(text, out, title_override=args.title)
    print(f"✅ 已生成: {out}")


if __name__ == "__main__":
    main()
