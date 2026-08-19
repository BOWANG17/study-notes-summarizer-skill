#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_docx.py — Render a "summary Markdown" into a formatted .docx (pure python-docx, no external deps).

This is the final step of the study-notes-summarizer skill pipeline (replacing WorkBuddy's
built-in tencent-docx skill), so rendering runs on any platform (Windows / macOS / Linux)
with zero WorkBuddy dependency.

Supported Markdown structures (aligned with references/section_guide.md and summary_prompt.md):
  # Heading                  -> document title (centered, enlarged)
  ## / ### Section heading    -> level-1 / level-2 sub-heading (six-section colors)
  | a | b | c |             -> vocab / comparison table (first row header, shaded)
  > ⚠️ Common mistake: ...   -> common-mistake card (light-red shading, "⚠️ Common mistake:" bold)
  > Covered files: ...       -> metadata box (light-blue shading)
  - / * list item           -> bullet list
  **bold**                   -> inline bold
  ---                        -> horizontal rule

Usage:
  python3 scripts/render_docx.py <input.md> -o <output.docx> [--title "Custom Title"]
"""

import argparse
import os
import platform
import re
import sys

# ---- bootstrap: auto-install python-docx if missing (matches parse_notes.py zero-manual strategy) ----
try:
    import docx
except ImportError:
    import subprocess
    print("📦 First run: auto-installing python-docx ...", file=sys.stderr)
    subprocess.run([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"], check=True)
    import docx

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------- colors / fonts ----------------
# Six-section heading colors (deep-blue family: distinct but not harsh)
SECTION_COLORS = {
    "Vocabulary": RGBColor(0x1F, 0x4E, 0x79),
    "Grammar":    RGBColor(0x2E, 0x6F, 0x4F),
    "Listening":  RGBColor(0x8A, 0x5A, 0x00),
    "Speaking":   RGBColor(0x9C, 0x27, 0x6B),
    "Reading":    RGBColor(0x37, 0x56, 0x9B),
    "Writing":    RGBColor(0x6A, 0x3D, 0x9A),
}
DEFAULT_SECTION_COLOR = RGBColor(0x1F, 0x4E, 0x79)

HEADER_FILL = "DCE6F1"     # table header light blue
TABLE_FILL  = "F2F6FB"     # table alternating-row light color (uniform light here)
WARN_FILL   = "FDE9E7"     # ⚠️ common-mistake card light red
INFO_FILL   = "E8F1FB"     # metadata box light blue
WARN_PREFIX = "⚠️ Common mistake"

# CJK font: pick a widely-available one per platform to avoid tofu (blank) boxes for CJK
def _cjk_font():
    if platform.system() == "Windows":
        return "Microsoft YaHei"
    if platform.system() == "Darwin":
        return "PingFang SC"
    return "Noto Sans CJK SC"


# ---------------- low-level helpers ----------------
def _set_cell_shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_para_shade(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_run_font(run, size=None, bold=None, color=None, font=None):
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if font:
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font)


def _add_inline(paragraph, text, base_size=11, base_font=None, base_color=None):
    """Parse **bold** inline markup and write runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size=base_size, bold=True, color=base_color, font=base_font)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size=base_size, color=base_color, font=base_font)


def _add_page_number_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_run_font(run, size=9, font=_cjk_font())
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    # PAGE field
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def _section_color(text):
    # take heading text after stripping '#', match by prefix/exact
    t = text.strip().lstrip("#").strip()
    for key, col in SECTION_COLORS.items():
        if t == key or t.startswith(key):
            return col
    return DEFAULT_SECTION_COLOR


# ---------------- inline parsing ----------------
def _strip_table_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def _is_table_sep(line):
    s = line.strip().replace("|", "").replace(" ", "")
    return bool(s) and set(s) <= set("-:")


# ---------------- main render ----------------
def render(md_text, out_path, title_override=None):
    doc = Document()
    font = _cjk_font()

    # default body font + size
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font)

    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    title_set = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # skip blank lines
        if not stripped:
            i += 1
            continue

        # heading H1
        if stripped.startswith("# ") and not title_set:
            title_text = stripped[2:].strip()
            if title_override:
                title_text = title_override
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            _set_run_font(run, size=20, bold=True, font=font, color=RGBColor(0x1F, 0x4E, 0x79))
            p.space_after = Pt(10)
            title_set = True
            i += 1
            continue

        # sub-section heading
        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:].strip())
            _set_run_font(run, size=13, bold=True, font=font, color=_section_color(stripped[4:]))
            p.space_before = Pt(8); p.space_after = Pt(4)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            _set_run_font(run, size=15, bold=True, font=font, color=_section_color(stripped[3:]))
            p.space_before = Pt(12); p.space_after = Pt(6)
            i += 1
            continue

        # horizontal rule
        if stripped == "---" or stripped == "***":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "BBBBBB")
            pbdr.append(bottom); pPr.append(pbdr)
            i += 1
            continue

        # blockquote (possibly multi-line)
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            quote_text = "\n".join(quote_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            if WARN_PREFIX in quote_text:
                _set_para_shade(p, WARN_FILL)
                # bold the "⚠️ Common mistake:" prefix, rest normal
                idx = quote_text.find(WARN_PREFIX)
                head = quote_text[:idx + len(WARN_PREFIX)]
                tail = quote_text[idx + len(WARN_PREFIX):]
                r1 = p.add_run(head)
                _set_run_font(r1, size=10.5, bold=True, font=font, color=RGBColor(0xB0, 0x30, 0x20))
                if tail:
                    _add_inline(p, tail, base_size=10.5, base_font=font)
            else:
                _set_para_shade(p, INFO_FILL)
                _add_inline(p, quote_text, base_size=10, base_font=font,
                            base_color=RGBColor(0x33, 0x55, 0x77))
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _strip_table_row(stripped)
            i += 2  # skip header + separator row
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_strip_table_row(lines[i].strip()))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for c, htext in enumerate(header):
                para = hdr[c].paragraphs[0]
                para.text = htext
                _set_run_font(para.runs[0], size=10.5, bold=True, font=font,
                              color=RGBColor(0x1F, 0x4E, 0x79))
                _set_cell_shade(hdr[c], HEADER_FILL)
            for r, row in enumerate(rows):
                cells = table.add_row().cells
                for c, ctext in enumerate(row):
                    cells[c].text = ""
                    _add_inline(cells[c].paragraphs[0], ctext, base_size=10.5, base_font=font)
                    if r % 2 == 1:
                        _set_cell_shade(cells[c], TABLE_FILL)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # bullet list
        if re.match(r"^[-*]\s+", stripped):
            items = []
            while i < n and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            for it in items:
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, it, base_size=11, base_font=font)
            continue

        # normal paragraph
        p = doc.add_paragraph()
        _add_inline(p, stripped, base_size=11, base_font=font)
        i += 1

    _add_page_number_footer(doc)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description="Markdown → formatted .docx (pure python-docx)")
    ap.add_argument("input", help="Input Markdown file")
    ap.add_argument("-o", "--output", required=True, help="Output .docx path")
    ap.add_argument("--title", help="Override document title (default: the Markdown # heading)")
    args = ap.parse_args()

    if not os.path.isfile(args.input):
        print(f"❌ Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    text = open(args.input, encoding="utf-8").read()
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    out = render(text, args.output, title_override=args.title)
    print(f"✅ Generated: {out}")


if __name__ == "__main__":
    main()
